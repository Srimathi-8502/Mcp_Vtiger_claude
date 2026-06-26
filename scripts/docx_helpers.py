"""Shared helpers for generating Word documentation."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullet(doc: Document, text: str, bold: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold:
        r = p.add_run(bold)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, hdr in enumerate(headers):
        t.rows[0].cells[i].text = hdr
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def title_page(doc: Document, title: str, subtitle: str) -> None:
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(11)
    doc.add_paragraph()


def document_map(doc: Document, parts: list[tuple[str, str]]) -> None:
    h(doc, "Document Contents", 1)
    table(doc, ["Part", "Source / Description"], parts)
    doc.add_page_break()


def _parse_table_row(line: str) -> list[str] | None:
    if not line.strip().startswith("|"):
        return None
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    if all(re.fullmatch(r"-+", c.replace(" ", "")) or c == "---" for c in cells):
        return None
    return cells


def add_markdown_file(
    doc: Document,
    path: Path,
    *,
    skip_title: bool = False,
    heading_level_offset: int = 0,
) -> None:
    """Convert project markdown (headings, tables, lists, code) into Word paragraphs."""
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    table_headers: list[str] | None = None
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_headers, table_rows
        if table_headers:
            table(doc, table_headers, table_rows)
        table_headers = None
        table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("```"):
            flush_table()
            in_code = True
            i += 1
            continue

        if stripped.startswith("|"):
            row = _parse_table_row(line)
            if row:
                if table_headers is None:
                    table_headers = row
                else:
                    table_rows.append(row)
            i += 1
            continue
        flush_table()

        if stripped in ("---", "***", "___"):
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            if skip_title and level == 1:
                i += 1
                continue
            h(doc, text, min(level + heading_level_offset, 4))
            i += 1
            continue

        if stripped.startswith("- [ ]") or stripped.startswith("- [x]"):
            bullet(doc, stripped[5:].strip())
            i += 1
            continue

        if stripped.startswith("- "):
            bullet(doc, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            numbered(doc, re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and len(stripped) > 2:
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip("*"))
            r.italic = True
            i += 1
            continue

        para(doc, stripped)
        i += 1

    flush_table()
    if in_code and code_lines:
        code_block(doc, "\n".join(code_lines))
