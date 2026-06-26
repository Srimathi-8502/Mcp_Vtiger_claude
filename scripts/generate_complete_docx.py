"""Generate Complete_MCP_Implementation.docx — full project documentation."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from docx_helpers import bullet, h, numbered, table

OUTPUT = Path(__file__).resolve().parent.parent / "Complete_MCP_Implementation.docx"


def add_stakeholder_sections(doc: Document, *, include_cover: bool = True) -> None:
    if include_cover:
        title = doc.add_heading("Complete MCP Implementation", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(
            "Uniware Vtiger CRM + Claude AM Copilot — Phase 1\n"
            "Document version: June 2026 | Author: Srimathi / Uniware GenAI Team"
        )
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(11)

    # ── 1. EXECUTIVE SUMMARY ──
    h(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "Phase 1 delivers a read-only Model Context Protocol (MCP) server that connects "
        "Anthropic Claude to Uniware's Vtiger CRM instance. Account Managers (AMs) can "
        "ask Claude in natural language for morning briefings — overdue follow-ups, open "
        "deals, and leads — without logging into Vtiger manually."
    )
    doc.add_paragraph(
        "The server is deployed on Render.com (free tier), monitored by UptimeRobot, "
        "and connected to Claude via a custom MCP connector. Source code is maintained "
        "on GitHub and GitLab."
    )
    table(doc, ["Item", "Value"], [
        ["GitHub Repository", "https://github.com/Srimathi-8502/Mcp_Vtiger_claude"],
        ["GitLab Repository", "http://192.168.1.195/genai/claude_vtiger_mcp"],
        ["Live Service URL", "https://mcp-vtiger-claude.onrender.com"],
        ["MCP Endpoint", "https://mcp-vtiger-claude.onrender.com/mcp"],
        ["Health Check", "https://mcp-vtiger-claude.onrender.com/health"],
        ["Vtiger Instance", "https://usp.od2.vtiger.com"],
        ["Phase", "Phase 1 — Read-only"],
    ])

    # ── 2. OBJECTIVE ──
    h(doc, "2. Project Objective", 1)
    bullet(doc, "Give Account Managers a daily CRM briefing through Claude conversational AI.")
    bullet(doc, "Eliminate manual Vtiger report building for morning standups.")
    bullet(doc, "Provide a secure, read-only integration — no CRM write access from AI.")
    bullet(doc, "Establish foundation for Phase 2 (write-back, user lookup, scheduling).")

    h(doc, "2.1 Success Criteria (Phase 1)", 2)
    bullet(doc, "Claude can call MCP tools and return real Vtiger data.")
    bullet(doc, "Deals show customer names (not raw Account IDs).")
    bullet(doc, "Overdue deal follow-ups are queryable per signed-in AM (no owner ID in prompt).")
    bullet(doc, "Service is deployed, health-checked, and demo-ready.")

    # ── 3. REQUIREMENTS ──
    h(doc, "3. Requirements", 1)

    h(doc, "3.1 Functional Requirements", 2)
    table(doc, ["ID", "Requirement", "Status"], [
        ["FR-01", "Fetch leads assigned to an AM by Vtiger user ID", "Done"],
        ["FR-02", "Fetch deals (Potentials) assigned to an AM", "Done"],
        ["FR-03", "Fetch overdue/due-today deal follow-ups", "Done"],
        ["FR-04", "Resolve deal organisation ID to Account name", "Done"],
        ["FR-05", "Expose data to Claude via MCP HTTP protocol", "Done"],
        ["FR-06", "Configurable CRM field mapping via environment", "Done"],
        ["FR-07", "Health and readiness endpoints", "Done"],
        ["FR-08", "Lead overdue follow-ups by date field", "Blocked — no Leads date field in Vtiger"],
        ["FR-09", "Per-user access via Microsoft Entra ID OAuth", "Done"],
        ["FR-10", "whoami tool to confirm signed-in AM and scope", "Done"],
        ["FR-11", "Server-side terminal stage exclusion", "Not in Phase 1 — prompt-side filter"],
        ["FR-12", "Server-side recency window filter", "Not in Phase 1 — prompt-side filter"],
        ["FR-13", "Write-back to Vtiger (activities, updates)", "Phase 2"],
    ])

    h(doc, "3.2 Non-Functional Requirements", 2)
    table(doc, ["ID", "Requirement", "Implementation"], [
        ["NFR-01", "Read-only CRM access", "No create/update/delete in codebase"],
        ["NFR-02", "Microsoft Entra ID OAuth authentication", "AzureProvider + email → Vtiger user mapping"],
        ["NFR-03", "Secrets not in git", ".env in .gitignore; Render env vars"],
        ["NFR-04", "HTTPS in production", "Render provides TLS"],
        ["NFR-05", "SQL injection protection", "Field name validation + value quoting"],
        ["NFR-06", "Paginated Vtiger queries", "LIMIT pagination up to 100/page"],
        ["NFR-07", "Availability for demo", "UptimeRobot pings /health every 5 min"],
    ])

    # ── 4. USE CASES ──
    h(doc, "4. Use Cases", 1)

    h(doc, "4.1 Primary Use Case — Morning AM Briefing", 2)
    doc.add_paragraph(
        "Actor: Account Manager\n"
        "Trigger: 9:00 AM daily standup or on-demand\n"
        "Flow: AM opens Claude → asks for briefing → Claude calls MCP tools → "
        "returns prioritized action list with deal names, customer names, stages, follow-up dates."
    )

    h(doc, "4.2 Example User Prompts", 2)
    for p in [
        '"Give me my morning CRM briefing."',
        '"What follow-ups are overdue for me today?"',
        '"Show my open deals with customer names and stages."',
        '"List all my leads and their status."',
        '"Who am I signed in as?" (calls whoami)',
    ]:
        bullet(doc, p)

    h(doc, "4.3 Who Benefits", 2)
    table(doc, ["Stakeholder", "Benefit"], [
        ["Account Managers", "Daily action list without opening Vtiger"],
        ["Sales Leadership", "Better follow-up discipline, fewer missed deals"],
        ["IT / Security", "Controlled read-only API bridge; no direct DB access"],
        ["GenAI Team", "Reusable MCP pattern for other Uniware integrations"],
    ])

    # ── 5. ARCHITECTURE ──
    h(doc, "5. System Architecture", 1)
    doc.add_paragraph(
        "Claude (claude.ai) → HTTP MCP → FastMCP Server (Python/Uvicorn on Render) "
        "→ Vtiger webservice.php → Vtiger CRM Cloud (usp.od2.vtiger.com)"
    )

    h(doc, "5.1 Component Stack", 2)
    table(doc, ["Layer", "Technology", "Purpose"], [
        ["AI Client", "Anthropic Claude", "Natural language interface for AMs"],
        ["Protocol", "MCP (Model Context Protocol)", "Tool calling standard"],
        ["Server Framework", "FastMCP 3.x + Starlette", "MCP HTTP app"],
        ["Runtime", "Python 3.12 + Uvicorn", "ASGI web server"],
        ["Hosting", "Render.com (Web Service)", "Cloud deployment"],
        ["Monitoring", "UptimeRobot", "Keep-alive + uptime checks"],
        ["CRM API", "Vtiger webservice.php", "Legacy REST-like API"],
        ["HTTP Client", "httpx (async)", "Vtiger API calls"],
        ["Config", "Pydantic Settings + .env", "Environment-based configuration"],
    ])

    h(doc, "5.2 Request Flow", 2)
    numbered(doc, "AM sends prompt in Claude with Vtiger_MCP connector enabled (Microsoft sign-in).")
    numbered(doc, "Claude calls MCP tool (e.g. get_my_overdue_followups) via HTTPS POST to /mcp.")
    numbered(doc, "Azure OAuth token validated; email mapped to Vtiger owner ID.")
    numbered(doc, "FastMCP routes to tool handler in tools/crm.py.")
    numbered(doc, "VtigerClient logs into Vtiger (getchallenge + MD5 accessKey).")
    numbered(doc, "Vtiger SQL-like query executed with pagination.")
    numbered(doc, "For deals: related_to Account IDs resolved to accountname.")
    numbered(doc, "JSON returned to Claude; Claude formats natural-language briefing.")

    h(doc, "5.3 Project Structure", 2)
    for line in [
        "src/vtiger_mcp/main.py — Entry point, health routes, FastMCP app",
        "src/vtiger_mcp/config.py — Environment settings and field mapping",
        "src/vtiger_mcp/middleware.py — API key validation",
        "src/vtiger_mcp/auth/users.py — OAuth identity, email → Vtiger owner mapping",
        "src/vtiger_mcp/tools/crm.py — Four MCP tool definitions (whoami + 3 briefing tools)",
        "src/vtiger_mcp/vtiger/client.py — Vtiger login, queries, normalization",
        "scripts/ — connection_check, describe_fields, describe_raw, test_tools",
        "docs/TECHNICAL.md — Architecture documentation",
        "render.yaml — Render Blueprint configuration",
        "Dockerfile — Container build (alternative deploy)",
    ]:
        bullet(doc, line)

    # ── 6. IMPLEMENTATION ──
    h(doc, "6. What Was Implemented", 1)

    h(doc, "6.1 MCP Tools (4)", 2)
    table(doc, ["Tool", "Purpose", "Vtiger Module / Scope"], [
        ["whoami", "Show signed-in AM email, Vtiger owner ID, admin flag, and data scope", "Users (lookup only)"],
        ["get_my_leads", "Leads for the signed-in AM (admin: all or filter by am_email)", "Leads"],
        ["get_my_deals", "Deals for the signed-in AM (admin: all or filter by am_email)", "Potentials"],
        ["get_my_overdue_followups", "Overdue deals/leads for signed-in AM (admin: all or filter)", "Leads + Potentials"],
    ])

    h(doc, "6.1.1 whoami Tool", 2)
    doc.add_paragraph(
        "The whoami tool confirms which Account Manager is authenticated via Microsoft Entra ID "
        "and what CRM data scope they have. AMs do not need to provide a Vtiger owner ID in the prompt — "
        "Claude or the AM can call whoami first to verify identity before running briefing tools."
    )
    table(doc, ["Output field", "Description", "Example"], [
        ["email", "Signed-in Microsoft / Uniware work email", "nirmal@uniware.net"],
        ["display_name", "AM name from Vtiger Users", "Nirmal Kumar J P"],
        ["owner_id", "Vtiger user ID mapped from email", "19x6"],
        ["is_admin", "True if email is in VTIGER_ADMIN_EMAILS", "true for Nirmal"],
        ["scope", "own_data_only or all_am_data", "all_am_data for admin"],
    ])
    doc.add_paragraph("Example whoami response:")
    doc.add_paragraph(
        '{"email": "nirmal@uniware.net", "display_name": "Nirmal Kumar J P", '
        '"owner_id": "19x6", "is_admin": true, "scope": "all_am_data"}'
    )

    h(doc, "6.2 Features Delivered", 2)
    bullet(doc, "Vtiger webservice login with MD5 challenge handshake")
    bullet(doc, "Paginated query_all() for large datasets")
    bullet(doc, "Configurable field mapping via VTIGER_FIELD_* env vars")
    bullet(doc, "Deal organisation_id → organisation_name resolution via Accounts module")
    bullet(doc, "Normalized JSON output keys for Claude (not raw Vtiger field names)")
    bullet(doc, "GET /health (liveness) and GET /ready (config check)")
    bullet(doc, "Microsoft Entra ID OAuth (AzureProvider) for per-user Claude access")
    bullet(doc, "whoami tool for signed-in AM identity and scope confirmation")
    bullet(doc, "get_my_* tools — no Vtiger owner ID required in prompts")
    bullet(doc, "Helper scripts for field discovery and connection testing")
    bullet(doc, "Render deployment with render.yaml Blueprint")
    bullet(doc, "GitHub + GitLab repository with README and technical docs")
    bullet(doc, "Word documentation (Vtiger_Claude_MCP.docx, this document)")

    h(doc, "6.3 HTTP Endpoints", 2)
    table(doc, ["Endpoint", "Auth", "Purpose"], [
        ["/health", "Public", "Liveness — always 200 if process up"],
        ["/ready", "Public", "Readiness — 200 if config complete, 503 if degraded"],
        ["/mcp", "OAuth (Microsoft)", "MCP protocol endpoint for Claude"],
    ])

    # ── 7. VTIGER DATA ──
    h(doc, "7. Vtiger Data — What We Pull", 1)

    h(doc, "7.1 Vtiger Modules Used", 2)
    table(doc, ["UI Name", "API Module Name", "Notes"], [
        ["Leads", "Leads", "describe confirmed"],
        ["Deals", "Potentials", "NOT 'Deals' — ACCESS_DENIED on describe"],
        ["Accounts", "Accounts", "Used to resolve organisation names"],
    ])

    h(doc, "7.2 Deals (Potentials) — Field Mapping", 2)
    table(doc, ["Briefing Field", "Vtiger API Name", "Env Variable"], [
        ["Assigned AM", "assigned_user_id", "VTIGER_FIELD_DEAL_OWNER"],
        ["Deal name", "potentialname", "VTIGER_FIELD_DEAL_NAME"],
        ["Organization (customer)", "related_to → accountname", "VTIGER_FIELD_DEAL_ORG"],
        ["Deal stage", "sales_stage", "VTIGER_FIELD_DEAL_STAGE"],
        ["Amount", "amount", "VTIGER_FIELD_DEAL_AMOUNT"],
        ["Last contacted", "last_contacted_on", "VTIGER_FIELD_DEAL_LAST_CONTACTED"],
        ["Next follow-up date", "cf_potentials_nextfollowupdate", "VTIGER_FIELD_DEAL_FOLLOWUP_DATE"],
        ["Next follow-up description", "nextstep", "VTIGER_FIELD_DEAL_FOLLOWUP_DESC"],
    ])

    h(doc, "7.3 Deal Sales Stages (Vtiger picklist)", 2)
    stages = (
        "New, Qualifying, Requirements Gathering, Demo POC, To Quote, Value Proposition, "
        "Negotiation, Ready to close, Monthy Billing, Mail confirmation, Execution pending, "
        "Closed Won, Closed Lost, Dormant, Payment Followup, Dropped, Inactive"
    )
    doc.add_paragraph(stages)

    h(doc, "7.4 Leads — Field Mapping", 2)
    table(doc, ["Briefing Field", "Vtiger API Name", "Env Variable"], [
        ["Assigned AM", "assigned_user_id", "VTIGER_FIELD_LEAD_OWNER"],
        ["Company", "company", "VTIGER_FIELD_LEAD_ORG"],
        ["Lead status", "leadstatus", "VTIGER_FIELD_LEAD_STATUS"],
        ["Next follow-up date", "NOT AVAILABLE in Vtiger Leads", "VTIGER_FIELD_LEAD_FOLLOWUP_DATE (empty)"],
        ["Notes / description", "description", "VTIGER_FIELD_LEAD_FOLLOWUP_DESC"],
    ])
    doc.add_paragraph(
        "Important: Vtiger Leads module has no 'Next Followup date' custom field (confirmed via "
        "describe operation June 2026). Therefore get_overdue_followups returns deals only for "
        "overdue filtering. Leads are included via get_my_leads."
    )

    h(doc, "7.5 MCP Output Keys (what Claude sees)", 2)
    h(doc, "whoami response", 3)
    for k in ["email, display_name, owner_id, is_admin, scope"]:
        bullet(doc, k)

    h(doc, "Deal record", 3)
    for k in [
        "id, owner, deal_name, organisation_id, organisation_name, stage, amount, "
        "last_contacted_date, next_followup_date, next_followup_description"
    ]:
        bullet(doc, k)

    h(doc, "Lead record", 3)
    for k in [
        "id, owner, organisation_name, status, next_followup_date, next_followup_description"
    ]:
        bullet(doc, k)

    # ── 8. BRIEFING ──
    h(doc, "8. What You Get in a Morning Briefing", 1)
    doc.add_paragraph(
        "When an AM asks Claude for a morning briefing, the following information "
        "can be assembled from MCP tool responses:"
    )
    table(doc, ["Briefing Section", "Source Tool", "Data Included"], [
        ["Signed-in AM identity", "whoami", "Email, Vtiger owner ID, admin flag, scope"],
        ["Overdue / due-today deals", "get_my_overdue_followups", "Deal name, customer name, stage, follow-up date, next step"],
        ["All open deals", "get_my_deals", "Deal name, customer, stage, amount, last contacted, follow-up"],
        ["All leads", "get_my_leads", "Company name, status, description/notes"],
        ["Prioritized action list", "Claude synthesis", "Claude ranks and formats for standup"],
    ])

    h(doc, "8.1 Filtering (Current State)", 2)
    bullet(doc, "Server-side open-stage whitelist: available via VTIGER_OPEN_DEAL_STAGES (currently empty = all stages).")
    bullet(doc, "Terminal stage exclusion (Closed Won/Lost etc.): NOT server-side — applied in Claude prompt.")
    bullet(doc, "Recency window (e.g. 90 days): NOT server-side — applied in Claude prompt.")
    bullet(doc, "Lead status exclusion (Disqualified etc.): NOT server-side — applied in Claude prompt.")

    h(doc, "8.2 Recommended Demo Prompt", 2)
    doc.add_paragraph(
        "You have access to the Uniware Vtiger MCP connector.\n\n"
        "STEP 1 — Call whoami to confirm who is signed in.\n"
        "STEP 2 — Call get_my_overdue_followups.\n"
        "STEP 3 — Call get_my_deals.\n"
        "STEP 4 — Call get_my_leads.\n\n"
        "Do not ask for a Vtiger owner ID — the server uses the signed-in Microsoft email.\n\n"
        "FILTERING (apply in summary): Exclude terminal deal stages (Closed Won, Closed Lost, "
        "Dormant, Dropped, Inactive). Use organisation_name not IDs. Exclude Disqualified leads.\n\n"
        "OUTPUT: Signed-in AM, overdue deals, open deals, active leads, Top 5 actions for today."
    )

    # ── 9. SECURITY ──
    h(doc, "9. Security", 1)

    h(doc, "9.1 Security Controls Implemented", 2)
    table(doc, ["Control", "Detail"], [
        ["Read-only", "No Vtiger write operations in any tool"],
        ["Microsoft Entra OAuth", "AzureProvider; email mapped to Vtiger owner ID"],
        ["Per-user data scope", "Regular AMs see own data; admins see all (VTIGER_ADMIN_EMAILS)"],
        ["Secrets in env only", ".env never committed; Render dashboard for production"],
        ["Field validation", "Module/field names restricted to [A-Za-z0-9_]"],
        ["SQL value escaping", "_quote() escapes single quotes in query values"],
        ["HTTPS", "Render provides TLS termination"],
        ["Vtiger session", "Per-process session cache; login via standard Vtiger challenge"],
    ])

    h(doc, "9.2 Security with Render", 2)
    bullet(doc, "Environment variables stored in Render dashboard (encrypted at rest by Render).")
    bullet(doc, "VTIGER_ACCESS_KEY and credentials never in source code.")
    bullet(doc, "Free tier: service sleeps on idle — reduces attack surface when not in use.")
    bullet(doc, "Public URL exposes only /health (no auth) and /mcp (API key if configured).")
    bullet(doc, "For demo: MCP_API_KEY was cleared to allow Claude connector (no-auth mode).")
    bullet(doc, "Risk: Without API key, anyone with /mcp URL can call tools. Re-enable MCP_API_KEY for production.")

    h(doc, "9.3 Security with Claude", 2)
    bullet(doc, "Claude custom connector uses OAuth (Microsoft work account) via Entra ID.")
    bullet(doc, "Leave OAuth Client ID/Secret empty in Claude — server holds Azure credentials.")
    bullet(doc, "Claude does not store Vtiger credentials — only MCP server has them.")
    bullet(doc, "AM does not provide Vtiger owner ID — whoami and get_my_* tools use signed-in email.")
    bullet(doc, "Claude Free plan may hit capacity limits unrelated to MCP security.")

    h(doc, "9.4 Security with UptimeRobot", 2)
    bullet(doc, "UptimeRobot only performs GET /health — no credentials sent.")
    bullet(doc, "Monitor URL is public health endpoint — no sensitive data exposed.")
    bullet(doc, "Does not access /mcp or Vtiger APIs.")
    bullet(doc, "502 incidents during cold start are availability issues, not security breaches.")

    h(doc, "9.5 Security Recommendations (Post-Demo)", 2)
    numbered(doc, "Re-enable MCP_API_KEY on Render only if not using OAuth.")
    numbered(doc, "Ensure Azure client secret rotation before expiry (recommend 12 months).")
    numbered(doc, "Upgrade Render to paid tier for always-on + better SLA.")
    numbered(doc, "Rotate Vtiger access key if exposed.")
    numbered(doc, "Use Claude Team/Enterprise for org-controlled connector access.")

    # ── 10. RENDER ──
    h(doc, "10. Render Deployment", 1)

    h(doc, "10.1 Service Configuration", 2)
    table(doc, ["Setting", "Value"], [
        ["Service name", "Mcp_Vtiger_claude"],
        ["Repository", "github.com/Srimathi-8502/Mcp_Vtiger_claude"],
        ["Branch", "main"],
        ["Runtime", "Python 3"],
        ["Build command", "pip install -r requirements.txt"],
        ["Start command", "PYTHONPATH=src uvicorn vtiger_mcp.main:app --host 0.0.0.0 --port $PORT"],
        ["Health check path", "/health"],
        ["Plan", "Free (spins down after ~15 min idle)"],
        ["Region", "Oregon (US West)"],
    ])

    h(doc, "10.2 Render Environment Variables", 2)
    table(doc, ["Variable", "Example / Value", "Required"], [
        ["VTIGER_BASE_URL", "https://usp.od2.vtiger.com", "Yes"],
        ["VTIGER_USERNAME", "dhana@uniware.net", "Yes"],
        ["VTIGER_ACCESS_KEY", "(secret)", "Yes"],
        ["VTIGER_FIELD_DEAL_FOLLOWUP_DATE", "cf_potentials_nextfollowupdate", "Yes"],
        ["VTIGER_FIELD_DEAL_FOLLOWUP_DESC", "nextstep", "Yes"],
        ["VTIGER_FIELD_DEAL_LAST_CONTACTED", "last_contacted_on", "Yes"],
        ["VTIGER_FIELD_DEAL_ORG", "related_to", "Yes"],
        ["VTIGER_FIELD_LEAD_FOLLOWUP_DESC", "description", "Optional"],
        ["VTIGER_FIELD_LEAD_FOLLOWUP_DATE", "(empty — no field in CRM)", "No"],
        ["MCP_PUBLIC_BASE_URL", "https://mcp-vtiger-claude.onrender.com", "Yes"],
        ["AZURE_TENANT_ID", "(from IT)", "Yes"],
        ["AZURE_CLIENT_ID", "(from IT)", "Yes"],
        ["AZURE_CLIENT_SECRET", "(secret)", "Yes"],
        ["AZURE_REQUIRED_SCOPES", "mcp-access", "Yes"],
        ["OAUTH_ALLOWED_DOMAIN", "uniware.net", "Yes"],
        ["VTIGER_ADMIN_EMAILS", "nirmal@uniware.net", "Yes"],
        ["MCP_API_KEY", "(empty when using OAuth)", "No"],
        ["FASTMCP_STATELESS_HTTP", "true", "Auto via render.yaml"],
        ["PYTHONPATH", "src", "Auto via render.yaml"],
    ])

    h(doc, "10.3 Deploy Steps", 2)
    numbered(doc, "Push code to GitHub main branch.")
    numbered(doc, "Render auto-deploys on push (or Manual Deploy).")
    numbered(doc, "Set environment variables in Render dashboard.")
    numbered(doc, "Verify https://mcp-vtiger-claude.onrender.com/health returns healthy.")
    numbered(doc, "Verify https://mcp-vtiger-claude.onrender.com/ready shows oauth_enabled: true.")

    h(doc, "10.4 Render Free Tier Limitations", 2)
    bullet(doc, "Service sleeps after ~15 minutes of no traffic.")
    bullet(doc, "Cold start takes 30–60 seconds — causes slow first MCP call.")
    bullet(doc, "502 Bad Gateway possible during wake-up.")
    bullet(doc, "0.1 CPU / 512 MB RAM — slow for large Vtiger queries.")
    bullet(doc, "Mitigation: UptimeRobot pings every 5 minutes to keep warm.")

    # ── 11. UPTIMEROBOT ──
    h(doc, "11. UptimeRobot Monitoring", 1)

    h(doc, "11.1 Configuration", 2)
    table(doc, ["Setting", "Value"], [
        ["Monitor type", "HTTP(s)"],
        ["URL", "https://mcp-vtiger-claude.onrender.com/health"],
        ["Interval", "5 minutes (free plan minimum)"],
        ["Expected response", "HTTP 200"],
        ["Timeout", "60 seconds (recommended)"],
    ])

    h(doc, "11.2 Why /health and not /ready", 2)
    bullet(doc, "/health always returns 200 — reliable Up status.")
    bullet(doc, "/ready returns 503 when MCP_API_KEY is empty — would show false Down.")

    h(doc, "11.3 Known Monitor Behaviour", 2)
    bullet(doc, "502 incidents during Render cold start are transient (not a real outage).")
    bullet(doc, "After warm-up, response time ~175ms.")
    bullet(doc, "Early uptime % may be low until service stays consistently warm.")
    bullet(doc, "Optional: cron-job.org at 8:55 AM as extra pre-briefing warm-up.")

    # ── 12. CLAUDE CONNECTION ──
    h(doc, "12. Claude MCP Connector Setup", 1)
    table(doc, ["Setting", "Value"], [
        ["Connector name", "Vtiger_MCP"],
        ["URL", "https://mcp-vtiger-claude.onrender.com/mcp"],
        ["OAuth Client ID", "Leave empty"],
        ["OAuth Client Secret", "Leave empty"],
    ])

    h(doc, "12.1 Connection Steps", 2)
    numbered(doc, "Claude → Settings → Connectors.")
    numbered(doc, "Add custom connector (or delete old failed one and recreate).")
    numbered(doc, "URL: https://mcp-vtiger-claude.onrender.com/mcp (must include /mcp).")
    numbered(doc, "Leave OAuth fields empty in Claude (server handles Microsoft sign-in).")
    numbered(doc, "On first use, sign in with @uniware.net Microsoft work account.")
    numbered(doc, "New chat → enable Vtiger_MCP → send briefing prompt.")

    h(doc, "12.2 Common Connection Errors", 2)
    table(doc, ["Error", "Cause", "Fix"], [
        ["Couldn't register sign-in service", "Claude expects OAuth; API key set on server", "Clear MCP_API_KEY on Render"],
        ["Couldn't reach MCP server", "Wrong URL (missing /mcp) or cold start", "Add /mcp; warm /health first"],
        ["502 Bad Gateway", "Render waking up", "Wait 60s; UptimeRobot keeps warm"],
        ["Capacity constraints", "Claude Free plan limit", "Retry later or upgrade Claude"],
        ["A bit longer...", "MCP tools running (normal)", "Wait; use shorter prompt"],
    ])

    # ── 13. LIMITATIONS ──
    h(doc, "13. Known Limitations (Phase 1)", 1)
    for item in [
        "Read-only — cannot log calls, update deals, or create records.",
        "AM identified by Microsoft email — no owner ID in prompt (whoami confirms identity).",
        "No lead overdue follow-up — Leads module lacks follow-up date field.",
        "No server-side terminal stage or recency filtering.",
        "Claude connector requires Microsoft OAuth setup in Azure Entra ID.",
        "Render free tier cold starts cause 30–120s delays on first request.",
        "Large deal/lead lists slow due to multiple Vtiger API round-trips.",
        "organisation_name resolution adds extra Accounts queries per deal batch.",
    ]:
        bullet(doc, item)

    # ── 14. PHASE 2 ──
    h(doc, "14. Phase 2 Roadmap", 1)
    table(doc, ["Feature", "Description"], [
        ["resolve_owner", "Lookup Vtiger user ID by email/name (partially done via whoami)"],
        ["create_activity", "Log call/meeting back to Vtiger"],
        ["update_followup", "Set next follow-up date from Claude"],
        ["Server-side stage filter", "Exclude terminal stages in get_deals_by_owner"],
        ["Server-side recency window", "Filter by last_contacted_on age"],
        ["Lead follow-up date", "When CRM adds field, wire VTIGER_FIELD_LEAD_FOLLOWUP_DATE"],
        ["OAuth for Claude", "Proper auth without disabling API key"],
        ["Render paid tier", "Always-on, faster CPU"],
        ["Scheduled briefing", "9 AM automated prompt via Claude Team/cron"],
        ["Caching", "Reduce repeated Vtiger login/query overhead"],
    ])

    # ── 15. TROUBLESHOOTING ──
    h(doc, "15. Troubleshooting Guide", 1)
    table(doc, ["Symptom", "Check", "Action"], [
        ["/health fails", "Render dashboard logs", "Redeploy; check build errors"],
        ["/ready 503", "missing_configuration list", "Set required env vars on Render"],
        ["Claude no tools", "Connector enabled in chat?", "Enable Vtiger_MCP; new chat"],
        ["Empty deals", "Not signed in or wrong email?", "Call whoami; verify Microsoft email matches Vtiger"],
        ["organisation_name null", "related_to empty on deal?", "Normal for deals without Account"],
        ["Slow response", "Cold start or many records?", "Warm /health; shorter prompt"],
        ["UptimeRobot Down", "502 during cold start?", "Increase timeout; wait for warm"],
    ])

    # ── 16. DELIVERABLES ──
    h(doc, "16. Deliverables Completed", 1)
    for item in [
        "MCP server source code (Python 3.12, FastMCP)",
        "Four MCP tools: whoami, get_my_leads, get_my_deals, get_my_overdue_followups",
        "Microsoft Entra ID OAuth and per-user AM access control",
        "Deal organisation name resolution (related_to → accountname)",
        "GitHub repository: github.com/Srimathi-8502/Mcp_Vtiger_claude",
        "GitLab repository: 192.168.1.195/genai/claude_vtiger_mcp",
        "Render deployment: mcp-vtiger-claude.onrender.com",
        "UptimeRobot monitor on /health every 5 min",
        "README.md, docs/TECHNICAL.md",
        "Vtiger_Claude_MCP.docx (summary doc)",
        "Complete_MCP_Implementation.docx (this document)",
        "Helper scripts: connection_check, describe_fields, describe_raw, test_tools",
        "Confirmed Vtiger field mappings for Potentials (Deals)",
        "Claude connector configured and tested",
    ]:
        bullet(doc, item)

    # ── 17. CONTACTS ──
    h(doc, "17. Key Contacts", 1)
    table(doc, ["Role", "Responsibility"], [
        ["Natasha", "Claude connector, AM Copilot rollout, CRM field mapping"],
        ["Dhana", "Vtiger API credentials and access"],
        ["Srimathi", "MCP development, Render deploy, documentation"],
        ["Divya / Leadership", "Demo stakeholder (June 23)"],
    ])

    doc.add_page_break()
    h(doc, "Appendix A — Full Environment Variable Reference", 1)
    vars_rows = [
        ["ENVIRONMENT", "production", "Runtime mode"],
        ["HOST", "0.0.0.0", "Bind address"],
        ["PORT", "8000 (Render sets $PORT)", "Listen port"],
        ["LOG_LEVEL", "INFO", "Logging verbosity"],
        ["FASTMCP_STATELESS_HTTP", "true", "Stateless MCP HTTP mode"],
        ["MCP_PUBLIC_BASE_URL", "https://mcp-vtiger-claude.onrender.com", "Public OAuth base URL"],
        ["AZURE_TENANT_ID", "(secret)", "Microsoft Entra tenant ID"],
        ["AZURE_CLIENT_ID", "(secret)", "Entra app client ID"],
        ["AZURE_CLIENT_SECRET", "(secret)", "Entra app client secret"],
        ["AZURE_REQUIRED_SCOPES", "mcp-access", "OAuth scope for MCP access"],
        ["OAUTH_ALLOWED_DOMAIN", "uniware.net", "Allowed sign-in email domain"],
        ["VTIGER_ADMIN_EMAILS", "nirmal@uniware.net", "Emails with full CRM access"],
        ["MCP_API_KEY", "(empty with OAuth)", "Legacy API key auth"],
        ["MCP_AUTH_TOKEN", "(optional)", "FastMCP Bearer auth"],
        ["VTIGER_BASE_URL", "https://usp.od2.vtiger.com", "Vtiger instance"],
        ["VTIGER_USERNAME", "API user email", "Vtiger login"],
        ["VTIGER_ACCESS_KEY", "(secret)", "Vtiger webservice key"],
        ["VTIGER_LEADS_MODULE", "Leads", "Leads module name"],
        ["VTIGER_DEALS_MODULE", "Potentials", "Deals module name"],
        ["VTIGER_ACCOUNTS_MODULE", "Accounts", "Accounts module name"],
        ["VTIGER_FIELD_ACCOUNT_NAME", "accountname", "Account name field"],
        ["VTIGER_FIELD_LEAD_OWNER", "assigned_user_id", "Lead owner field"],
        ["VTIGER_FIELD_LEAD_ORG", "company", "Lead company field"],
        ["VTIGER_FIELD_LEAD_STATUS", "leadstatus", "Lead status field"],
        ["VTIGER_FIELD_LEAD_FOLLOWUP_DATE", "(empty)", "No field in CRM"],
        ["VTIGER_FIELD_LEAD_FOLLOWUP_DESC", "description", "Lead notes field"],
        ["VTIGER_FIELD_DEAL_OWNER", "assigned_user_id", "Deal owner field"],
        ["VTIGER_FIELD_DEAL_NAME", "potentialname", "Deal name field"],
        ["VTIGER_FIELD_DEAL_ORG", "related_to", "Deal organisation ref"],
        ["VTIGER_FIELD_DEAL_STAGE", "sales_stage", "Deal stage field"],
        ["VTIGER_FIELD_DEAL_AMOUNT", "amount", "Deal amount field"],
        ["VTIGER_FIELD_DEAL_LAST_CONTACTED", "last_contacted_on", "Last contacted"],
        ["VTIGER_FIELD_DEAL_FOLLOWUP_DATE", "cf_potentials_nextfollowupdate", "Follow-up date"],
        ["VTIGER_FIELD_DEAL_FOLLOWUP_DESC", "nextstep", "Follow-up description"],
        ["VTIGER_OPEN_DEAL_STAGES", "(empty)", "Optional stage whitelist"],
    ]
    table(doc, ["Variable", "Value", "Description"], vars_rows)

    h(doc, "Appendix B — Git Commit History (Key Milestones)", 1)
    table(doc, ["Commit", "Description"], [
        ["49b69e6", "Initial commit: Vtiger MCP server"],
        ["8396173", "Add README with setup and deployment docs"],
        ["2bca5e4", "Deal org mapping, technical docs, field config"],
        ["f901f6e", "Resolve related_to to Account name for briefings"],
        ["48c092d", "Azure OAuth and per-user AM access without owner IDs"],
        ["acb8e42", "Document Leads follow-up limits"],
    ])


def build() -> Document:
    doc = Document()
    add_stakeholder_sections(doc)
    return doc


def main() -> None:
    doc = build()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
