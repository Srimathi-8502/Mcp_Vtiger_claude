"""Generate Vtiger_Claude_MCP.docx project summary document."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUTPUT = Path(__file__).resolve().parent.parent / "Vtiger_Claude_MCP.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def build_document() -> Document:
    doc = Document()

    title = doc.add_heading("Uniware Vtiger + Claude MCP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Phase 1 — Project Summary, Use Cases & Next Steps")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)

    doc.add_paragraph(
        "Repository: http://192.168.1.195/genai/claude_vtiger_mcp\n"
        "Deployment platform: Render.com\n"
        "Document version: June 2026"
    )

    # --- 1. USE CASE ---
    add_heading(doc, "1. Business Use Case (Start Here)", 1)
    doc.add_paragraph(
        "Uniware Account Managers (AMs) manage leads and deals in Vtiger CRM. "
        "Each day they need a clear picture of their pipeline: which leads are active, "
        "which deals are open, and which records have follow-ups due or overdue."
    )
    doc.add_paragraph(
        "Today this requires logging into Vtiger, running reports, and manually "
        "building a daily action list. The AM Copilot vision is to let an AM ask "
        "Claude in natural language — for example:"
    )
    for example in [
        '"What follow-ups are overdue for me today?"',
        '"Show my open deals and their stages."',
        '"List all my leads and their status."',
        '"Prepare my morning briefing for user 19x5."',
    ]:
        add_bullet(doc, example)

    doc.add_paragraph(
        "Claude does not connect to Vtiger directly. Instead, Claude calls a "
        "Model Context Protocol (MCP) server that we built. This server translates "
        "Claude's tool requests into Vtiger webservice queries and returns structured "
        "JSON that Claude turns into a readable briefing."
    )

    add_heading(doc, "1.1 Who Benefits", 2)
    add_bullet(doc, " — Gets daily briefings without opening Vtiger", "Account Managers")
    add_bullet(doc, " — Faster follow-up discipline, fewer missed opportunities", "Sales leadership")
    add_bullet(doc, " — Read-only, API-key protected bridge between AI and CRM", "IT / Security")

    add_heading(doc, "1.2 Phase 1 Scope (What We Built)", 2)
    doc.add_paragraph("Phase 1 is intentionally read-only. Users can query data, not change it.")
    add_bullet(doc, "View leads assigned to an AM")
    add_bullet(doc, "View open deals assigned to an AM")
    add_bullet(doc, "Get overdue / due-today follow-ups across leads and deals")
    add_bullet(doc, "NOT in Phase 1: create leads, update deals, log calls, send emails")

    # --- 2. WHAT USERS CAN DO ---
    add_heading(doc, "2. What Users Can Do With This MCP", 1)

    add_heading(doc, "2.1 MCP Tools Available to Claude", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Tool Name"
    hdr[1].text = "What It Does"
    hdr[2].text = "Example Claude Prompt"
    rows = [
        (
            "get_leads_by_owner",
            "Returns all leads for a Vtiger user ID: organisation name, status, follow-up date & description.",
            '"List all leads owned by 19x5"',
        ),
        (
            "get_deals_by_owner",
            "Returns open deals (Potentials): name, stage, amount, last contacted, follow-up fields.",
            '"Show open deals for AM 19x5"',
        ),
        (
            "get_overdue_followups",
            "Primary briefing tool. Returns leads and deals where next follow-up date is today or earlier.",
            '"What follow-ups are overdue for 19x5 today?"',
        ),
    ]
    for name, desc, prompt in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = desc
        row[2].text = prompt

    doc.add_paragraph()

    add_heading(doc, "2.2 Typical User Conversations", 2)
    doc.add_paragraph("Morning briefing workflow:")
    add_numbered(doc, "AM opens Claude (with Vtiger MCP connector enabled).")
    add_numbered(doc, 'AM asks: "Give me my overdue follow-ups for today. My Vtiger ID is 19x5."')
    add_numbered(doc, "Claude calls get_overdue_followups(owner='19x5').")
    add_numbered(doc, "MCP server queries Vtiger, normalizes records, returns JSON.")
    add_numbered(doc, "Claude formats a prioritized action list in plain English.")
    add_numbered(doc, "AM works through the list in Vtiger (updates still done in CRM manually).")

    add_heading(doc, "2.3 Data Returned Per Record", 2)
    doc.add_paragraph("Leads:")
    for f in ["id", "owner", "organisation_name", "status", "next_followup_date", "next_followup_description"]:
        add_bullet(doc, f)
    doc.add_paragraph("Deals:")
    for f in [
        "id", "owner", "deal_name", "stage", "amount",
        "last_contacted_date", "next_followup_date", "next_followup_description",
    ]:
        add_bullet(doc, f)

    # --- 3. WHAT WE BUILT ---
    add_heading(doc, "3. What We Built — Technical Summary", 1)

    add_heading(doc, "3.1 Architecture", 2)
    doc.add_paragraph(
        "Claude Connector  →  HTTP (MCP)  →  FastMCP Server (Python/Uvicorn)  →  "
        "Vtiger webservice.php  →  Vtiger CRM Cloud"
    )

    add_heading(doc, "3.2 Components Completed", 2)
    components = [
        ("main.py", "Application entry point. Creates FastMCP server, registers tools, /health and /ready routes, exposes ASGI app for Uvicorn."),
        ("config.py", "Pydantic Settings loading all environment variables and CRM field mappings from .env."),
        ("middleware.py", "ApiKeyMiddleware — validates X-API-Key or Bearer token on MCP paths."),
        ("tools/crm.py", "Three MCP tools: get_leads_by_owner, get_deals_by_owner, get_overdue_followups."),
        ("vtiger/client.py", "Vtiger webservice client: login (MD5 challenge), paginated query, field normalization, SQL injection guards."),
        ("logging_config.py", "Structured stdout logging."),
        ("scripts/", "connection_check.py, mcp_check.py, describe_fields.py, test_tools.py for verification."),
        ("Dockerfile", "Container image for alternative deployment."),
        ("render.yaml", "Render.com Blueprint — build command, start command, health check, env var placeholders."),
        ("README.md", "Setup and run instructions."),
        ("docs/TECHNICAL.md", "Architecture and request-flow documentation."),
    ]
    for name, desc in components:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(name)
        r.bold = True
        p.add_run(f" — {desc}")

    add_heading(doc, "3.3 Request Flow (How Code Runs)", 2)
    flow_steps = [
        "Process starts via vtiger-mcp CLI, uvicorn, or Docker CMD.",
        "main.py loads settings from .env, configures logging, creates FastMCP instance.",
        "register_tools() binds three tool handlers; each shares one VtigerClient.",
        "app = mcp.http_app() wraps everything in Starlette ASGI with ApiKeyMiddleware.",
        "Incoming HTTP request hits middleware (auth check unless /health or /ready).",
        "FastMCP routes MCP tool call to the matching async function in crm.py.",
        "crm.py calls VtigerClient method → builds vtiger SQL query → query_all() with pagination.",
        "VtigerClient logs in (getchallenge + MD5 accessKey), runs query, normalizes field names.",
        "JSON string returned to Claude via MCP protocol.",
    ]
    for i, step in enumerate(flow_steps, 1):
        add_numbered(doc, step)

    add_heading(doc, "3.4 Security Features", 2)
    add_bullet(doc, "Read-only — no write operations to Vtiger")
    add_bullet(doc, "MCP_API_KEY required on all MCP endpoints in production")
    add_bullet(doc, "Field/module names validated (alphanumeric + underscore only)")
    add_bullet(doc, "User values SQL-escaped before query construction")
    add_bullet(doc, "Secrets in environment only — never committed to Git")

    # --- 4. NEXT STEPS ---
    add_heading(doc, "4. Next Steps — Connect Claude to MCP", 1)

    add_heading(doc, "4.1 Prerequisites Checklist", 2)
    checklist = [
        "Vtiger API credentials from CRM admin (username + access key)",
        "Correct VTIGER_BASE_URL confirmed (e.g. https://usp.od2.vtiger.com)",
        "CRM field API names confirmed for follow-up dates (cf_* custom fields)",
        "Vtiger user IDs for each AM who will use the copilot (e.g. 19x5)",
        "MCP_API_KEY generated (python -c \"import secrets; print(secrets.token_urlsafe(32))\")",
        "Server deployed and reachable over HTTPS (Render provides this)",
    ]
    for item in checklist:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ ").bold = True
        p.add_run(item)

    add_heading(doc, "4.2 Step-by-Step: Deploy on Render", 2)
    render_steps = [
        "Create account at https://render.com and connect your GitLab repository.",
        "In Render Dashboard → New → Blueprint (or Web Service).",
        "Point to repo: http://192.168.1.195/genai/claude_vtiger_mcp (or mirror on GitHub if Render requires it).",
        "Render reads render.yaml automatically if using Blueprint.",
        "Service settings from render.yaml:\n"
        "   • Build: pip install -r requirements.txt\n"
        "   • Start: PYTHONPATH=src uvicorn vtiger_mcp.main:app --host 0.0.0.0 --port $PORT\n"
        "   • Health check path: /health",
        "In Render → Environment, set all secret variables (sync: false in render.yaml):",
    ]
    for step in render_steps:
        add_numbered(doc, step)

    env_table = doc.add_table(rows=1, cols=2)
    env_table.style = "Table Grid"
    env_table.rows[0].cells[0].text = "Variable"
    env_table.rows[0].cells[1].text = "Value"
    env_vars = [
        ("MCP_API_KEY", "Your generated secret key"),
        ("VTIGER_BASE_URL", "https://usp.od2.vtiger.com (confirm with admin)"),
        ("VTIGER_USERNAME", "api.user@uniware.com"),
        ("VTIGER_ACCESS_KEY", "From Vtiger My Preferences → Access Key"),
        ("VTIGER_FIELD_DEAL_FOLLOWUP_DATE", "cf_potentials_nextfollowupdate (confirm)"),
        ("VTIGER_FIELD_DEAL_LAST_CONTACTED", "last_contacted_on"),
        ("VTIGER_FIELD_DEAL_FOLLOWUP_DESC", "nextstep"),
        ("VTIGER_OPEN_DEAL_STAGES", "Optional: Prospecting,Qualification,... (comma-separated)"),
    ]
    for var, val in env_vars:
        row = env_table.add_row().cells
        row[0].text = var
        row[1].text = val

    doc.add_paragraph()
    post_deploy = [
        "Deploy the service. Wait for build to complete (typically 2–5 minutes).",
        "Note your Render URL: https://uniware-vtiger-mcp.onrender.com (name may vary).",
        "Verify deployment:\n   GET https://<your-url>/health  → {\"status\": \"healthy\"}\n   GET https://<your-url>/ready   → {\"status\": \"ready\"} (503 if config missing)",
        "Run connection_check.py locally against production URL if needed.",
    ]
    for step in post_deploy:
        add_numbered(doc, step)

    add_heading(doc, "4.3 Step-by-Step: Register MCP in Claude", 2)
    claude_steps = [
        "Open Claude (claude.ai) → Settings → Connectors (or Custom Integrations / MCP).",
        "Add a new custom MCP connector / remote MCP server.",
        "Server URL: https://<your-render-url>/mcp",
        "Authentication: API Key header\n   Header name: X-API-Key\n   Value: <your MCP_API_KEY>",
        "Save and enable the connector for your workspace or project.",
        "Start a new Claude conversation with the connector enabled.",
        "Test: \"Call get_overdue_followups for owner 19x5\" or ask in natural language.",
        "Confirm Claude lists tools: get_leads_by_owner, get_deals_by_owner, get_overdue_followups.",
    ]
    for step in claude_steps:
        add_numbered(doc, step)

    doc.add_paragraph(
        "Note: Exact Claude UI labels may vary by plan (Team, Enterprise) and region. "
        "If your org uses Anthropic Admin console, Natasha / IT may register the connector centrally."
    )

    add_heading(doc, "4.4 Field Mapping Session (Required Before Go-Live)", 2)
    doc.add_paragraph(
        "Several follow-up fields are custom in Vtiger (cf_*). Run describe_fields.py "
        "or use Vtiger Settings → Studio to confirm exact API names, then update Render env vars:"
    )
    add_bullet(doc, "VTIGER_FIELD_LEAD_FOLLOWUP_DATE")
    add_bullet(doc, "VTIGER_FIELD_LEAD_FOLLOWUP_DESC")
    add_bullet(doc, "VTIGER_FIELD_DEAL_FOLLOWUP_DATE")
    add_bullet(doc, "VTIGER_FIELD_DEAL_FOLLOWUP_DESC")
    doc.add_paragraph(
        "Until at least one follow-up date field is set, /ready returns 503 (degraded). "
        "get_overdue_followups will only query modules whose date field is configured."
    )

    add_heading(doc, "4.5 Recommended Rollout Plan", 2)
    rollout = [
        "Week 1 — Deploy to Render, verify /health and /ready, run scripts against staging Vtiger.",
        "Week 1 — Field mapping session with Natasha / CRM admin; update env vars.",
        "Week 2 — Register MCP connector in Claude for pilot group (2–3 AMs).",
        "Week 2 — Pilot AMs test morning briefing prompts with their Vtiger user IDs.",
        "Week 3 — Gather feedback; adjust open deal stages filter and field mappings.",
        "Week 4+ — Wider rollout; plan Phase 2 (write operations, user lookup, activities).",
    ]
    for item in rollout:
        add_bullet(doc, item)

    # --- 5. PHASE 2 ---
    add_heading(doc, "5. Future Phase 2 (Not Built Yet)", 1)
    future = [
        "resolve_owner — lookup Vtiger user ID by email/name (AMs won't need to know 19x5)",
        "create_activity — log a call or meeting back to Vtiger",
        "update_followup — set next follow-up date and notes from Claude",
        "get_deal_pipeline_summary — aggregated metrics by stage",
        "OAuth / SSO instead of static API key",
        "Caching layer for high-volume queries",
    ]
    for item in future:
        add_bullet(doc, item)

    # --- 6. TROUBLESHOOTING ---
    add_heading(doc, "6. Troubleshooting", 1)
    troubles = [
        ("/ready returns 503", "Missing env vars. Check Render Environment tab for MCP_API_KEY, Vtiger credentials, and at least one follow-up date field."),
        ("Claude gets 401", "MCP_API_KEY mismatch between Render and Claude connector config."),
        ("Vtiger login fails", "Wrong base URL, username, or access key. Run connection_check.py."),
        ("Empty follow-up results", "Follow-up date field name wrong, or no records due. Verify with describe_fields.py."),
        ("Render free tier sleeps", "First request after idle may take 30–60s. Upgrade plan or use uptime ping."),
        ("Claude cannot see tools", "Connector URL must end at /mcp; connector must be enabled for the chat."),
    ]
    for problem, fix in troubles:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(problem + ": ")
        r.bold = True
        p.add_run(fix)

    # --- 7. CONTACTS ---
    add_heading(doc, "7. Key Contacts & References", 1)
    add_bullet(doc, " — Vtiger credentials, field names, user IDs", "Natasha / CRM admin (Dhana)")
    add_bullet(doc, " — Claude connector registration, AM Copilot rollout", "Natasha")
    add_bullet(doc, " — GitLab repo, Render deployment, code changes", "Development team")
    doc.add_paragraph()
    doc.add_paragraph("Documentation in repository:")
    add_bullet(doc, "README.md — setup and run")
    add_bullet(doc, "docs/TECHNICAL.md — architecture and code flow")
    add_bullet(doc, ".env.example — all configuration variables")

    doc.add_page_break()
    add_heading(doc, "Appendix A — Environment Variables (Full List)", 1)
    appendix_vars = [
        ("ENVIRONMENT", "production / development", "No"),
        ("HOST / PORT", "0.0.0.0 / 8000", "No (Render sets PORT)"),
        ("LOG_LEVEL", "INFO", "No"),
        ("FASTMCP_STATELESS_HTTP", "true", "No"),
        ("MCP_API_KEY", "Secret for Claude → server auth", "Yes"),
        ("MCP_AUTH_TOKEN", "Optional Bearer token", "No"),
        ("VTIGER_BASE_URL", "Vtiger instance URL", "Yes"),
        ("VTIGER_USERNAME", "API user email", "Yes"),
        ("VTIGER_ACCESS_KEY", "Vtiger access key", "Yes"),
        ("VTIGER_FIELD_*", "CRM field API name mappings", "Partial"),
        ("VTIGER_OPEN_DEAL_STAGES", "Comma-separated open stages", "No"),
    ]
    apt = doc.add_table(rows=1, cols=3)
    apt.style = "Table Grid"
    apt.rows[0].cells[0].text = "Variable"
    apt.rows[0].cells[1].text = "Description"
    apt.rows[0].cells[2].text = "Required"
    for var, desc, req in appendix_vars:
        row = apt.add_row().cells
        row[0].text = var
        row[1].text = desc
        row[2].text = req

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
